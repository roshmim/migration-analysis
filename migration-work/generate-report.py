#!/usr/bin/env python3
"""Generate Xero.com EDS Migration Analysis Word Document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

for level in range(1, 5):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.color.rgb = RGBColor(0x13, 0xB5, 0xEA)

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], '13B5EA')
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.bold = True
                r.font.size = Pt(9)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def add_screenshot(doc, path, caption, width=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.style.font.size = Pt(8)
        if cap.runs:
            cap.runs[0].italic = True

# TITLE PAGE
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Xero.com (US)\nEdge Delivery Services\nMigration Analysis Report')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x13, 0xB5, 0xEA)
run.bold = True

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Comprehensive Site Analysis & Migration Estimation')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('Date: April 24, 2026')
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()
summary_box = doc.add_paragraph()
summary_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = summary_box.add_run('Total Pages: ~313  |  Templates: 23  |  Blocks: 24  |  Integrations: 27')
run.font.size = Pt(11)
run.bold = True

doc.add_page_break()

# TABLE OF CONTENTS
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Templates Inventory',
    '2. Blocks / Components Catalog',
    '3. Page Counts by Template',
    '4. Integrations Analysis',
    '5. Complex Use Cases & Observations',
    '6. Migration Estimates',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.style.font.size = Pt(11)
doc.add_page_break()

# SECTION 1: TEMPLATES INVENTORY
doc.add_heading('1. Templates Inventory', level=1)
doc.add_paragraph('The following unique page templates were identified across the Xero.com US site. Each template represents a distinct page layout pattern that can be built once and reused across multiple pages.')

templates_data = [
    ['1', 'Homepage', 'High', 'Most complex page with hero, trust badges, AI section, feature cards, promotional grid, pricing table with toggle, business types carousel, partner CTA', 'https://www.xero.com/us/'],
    ['2', 'Product Feature Page', 'Medium-High', 'Feature detail pages with hero + CTAs, benefit cards, alternating image-text sections, testimonial, pricing table, FAQ accordion, related features, video embed', 'https://www.xero.com/us/accounting-software/send-invoices/'],
    ['3', 'Business Type Page', 'Medium', 'Industry-specific pages with hero, feature highlights, testimonial, pricing table, and CTAs. Templated layout reused across 46 industries', 'https://www.xero.com/us/small-businesses/construction/'],
    ['4', 'Business Type Listing', 'Low-Medium', 'Grid directory page listing all business types as linked cards with heading and description', 'https://www.xero.com/us/small-businesses/all-business-types/'],
    ['5', 'Accountant/Bookkeeper Page', 'Medium', 'Professional audience pages with hero, product features, partner program info, and practice tool CTAs', 'https://www.xero.com/us/accountants-bookkeepers/'],
    ['6', 'Practice Product Page', 'Medium', 'Product pages for practice tools (HQ, Practice Manager, Workpapers, Ledger/Cashbook, Syft)', 'https://www.xero.com/us/xero-hq/'],
    ['7', 'Pricing Page', 'High', 'Plan comparison with radio toggle (Buy Now / Free Trial), 3 plan cards with pricing, feature checklists, compare link', 'https://www.xero.com/us/pricing-plans/'],
    ['8', 'Guide Listing Page', 'Low-Medium', 'Hub page listing all guides with category cards and links in a grid', 'https://www.xero.com/us/guides/'],
    ['9', 'Guide Article Page', 'Medium', 'Long-form educational content with table of contents, inline CTAs, related guides, and structured article body', 'https://www.xero.com/us/guides/invoicing/'],
    ['10', 'Template Listing Page', 'Low-Medium', 'Hub listing downloadable business templates with category cards', 'https://www.xero.com/us/templates/'],
    ['11', 'Template Detail Page', 'Medium', 'Template preview with download CTA, description, step-by-step guide, and related templates', 'https://www.xero.com/us/templates/invoice-template/'],
    ['12', 'Glossary Listing Page', 'Low', 'Alphabetical A-Z index page with grouped term links', 'https://www.xero.com/us/glossary/'],
    ['13', 'Glossary Term Page', 'Low', 'Term definition with explanation, related terms, and CTA. 115 pages share this template', 'https://www.xero.com/us/glossary/cash-flow/'],
    ['14', 'Calculator Listing Page', 'Low', 'Hub listing all calculators with description cards', 'https://www.xero.com/us/calculators/'],
    ['15', 'Calculator Tool Page', 'Very High', 'Interactive calculator with form inputs, real-time JS computation, result display, FAQ accordion, and explanatory content', 'https://www.xero.com/us/calculators/margin-calculator/'],
    ['16', 'Compare Page', 'Medium-High', 'Competitor comparison with feature comparison table, benefit sections, testimonials, and switch CTAs', 'https://www.xero.com/us/versus/quickbooks-alternative/'],
    ['17', 'Company Info Page', 'Medium', 'Corporate pages (About, Contact, Sustainability, Investors, etc.) with varied content-heavy layouts', 'https://www.xero.com/us/about/'],
    ['18', 'Legal Page', 'Low', 'Text-heavy legal content (Terms, Privacy, Cookies) with standard heading/paragraph structure', 'https://www.xero.com/us/legal/terms/'],
    ['19', 'Campaign Page', 'Medium-High', 'Marketing landing pages with promotional content, forms, and CTAs', 'https://www.xero.com/us/affiliate-program/'],
    ['20', 'Resource Page', 'Medium', 'Resource hubs for insights, customer stories, and certifications with card grids', 'https://www.xero.com/us/resources/'],
    ['21', 'AI Feature Page', 'High', 'JAX AI superagent showcase with demos, feature highlights, and interactive elements', 'https://www.xero.com/us/ai-in-accounting/jax-vision/'],
    ['22', 'Utility Page', 'Low-Medium', 'Functional pages (Sitemap, Advisor Finder, Login, Data Protection, Integrators directory)', 'https://www.xero.com/us/sitemap/'],
    ['23', 'Signup Page', 'High', 'Multi-step registration wizard with form fields, reCAPTCHA, Salesforce CRM integration', 'https://www.xero.com/us/signup/early/'],
]

create_table(doc, ['#', 'Template Name', 'Complexity', 'Reasoning', 'Reference URL'], templates_data)

doc.add_paragraph()
add_screenshot(doc, '/tmp/playwright/02-homepage-full.png', 'Figure 1: Homepage Template - Full Page View', 3.5)
doc.add_paragraph()
add_screenshot(doc, '/tmp/playwright/03-product-feature-page.png', 'Figure 2: Product Feature Page Template (Send Invoices)', 3.5)
doc.add_paragraph()
add_screenshot(doc, '/tmp/playwright/04-pricing-page.png', 'Figure 3: Pricing Page Template', 3.5)

doc.add_page_break()

# SECTION 2: BLOCKS / COMPONENTS CATALOG
doc.add_heading('2. Blocks / Components Catalog', level=1)
doc.add_paragraph('The following reusable blocks and components were identified across the site. Design variations of the same content model are tracked as variants of a single block rather than separate blocks.')

doc.add_heading('2.1 Core Blocks (Extend from EDS Boilerplate)', level=2)

core_blocks = [
    ['1', 'header', 'Medium-High', 'Global navigation with promo banner, logo, mega-menu (desktop), hamburger menu (mobile), CTA, and Log in link. Promo banner includes countdown timer and dismiss functionality.', 'All pages'],
    ['2', 'footer', 'Medium', '6-column mega footer with categorized links, region selector, social media icons, legal links, and copyright.', 'All pages'],
    ['3', 'hero', 'Medium', 'Multiple variants: (a) Split hero with animated text + product image; (b) Feature hero with heading + CTA buttons + screenshot; (c) Simple heading hero with breadcrumb.', 'Homepage, Product pages, Business type pages'],
    ['4', 'cards', 'Medium', 'Multiple variants: (a) Icon + heading + text feature cards (3-up); (b) Image + heading + text linked cards; (c) Trust badge logo cards; (d) Business type cards.', 'Homepage, Feature pages, Business types, Resources'],
    ['5', 'columns', 'Low-Medium', 'Two-column image + text layout. Variants: (a) Image left / text right; (b) Text left / image right. For alternating feature sections.', 'Product pages, Business type pages, Campaign pages'],
    ['6', 'fragment', 'Low', 'Shared content fragments for reusable sections (pricing table, CTA banners) across multiple templates.', 'Multiple pages'],
]

create_table(doc, ['#', 'Block', 'Complexity', 'Description & Behaviour', 'Reference URL(s)'], core_blocks)

doc.add_heading('2.2 Custom Blocks (New Development Required)', level=2)

custom_blocks = [
    ['7', 'promo-banner', 'High', 'Dismissable top-of-page banner with countdown timer (days/hours/mins/secs), offer text, CTA. JS for timer logic, dismiss cookie persistence, A/B variant support.', 'All pages'],
    ['8', 'pricing-table', 'High', '3 plan cards with radio toggle (Buy Now / Free Trial), per-plan pricing with savings, feature bullets, CTAs. Responsive horizontal scroll on mobile.', 'Homepage, Product pages, Pricing pages'],
    ['9', 'pricing-detail', 'Medium', 'Individual plan detail page with full feature comparison checklist, plan pricing, and signup CTA.', 'Pricing detail pages'],
    ['10', 'tabs', 'Medium', 'Tab/radio toggle to switch between content panels. Used for pricing plan categories and feature sections.', 'Pricing pages, Feature pages'],
    ['11', 'comparison-table', 'Medium-High', 'Side-by-side feature comparison table with checkmarks, cross-marks, feature rows comparing Xero vs competitor.', 'Compare pages (vs QuickBooks, FreshBooks, Sage)'],
    ['12', 'accordion', 'Medium', 'Expandable FAQ sections with question heading + answer content. Supports multiple items.', 'Product pages, Calculator pages'],
    ['13', 'carousel', 'Medium', 'Horizontal scrolling component with navigation arrows for plan cards on mobile and business type cards.', 'Homepage'],
    ['14', 'feature-list', 'Low-Medium', 'Checkmark icon + text list for feature highlights and plan inclusions.', 'Product pages, Pricing pages, Compare pages'],
    ['15', 'trust-badges', 'Low', 'Row of award/recognition logos (G2, Capterra, TrustRadius, etc.) with alt text.', 'Homepage'],
    ['16', 'cta-banner', 'Low-Medium', 'Full-width promotional banner with heading, description, and action buttons. Variants: teal/green background.', 'All templates'],
    ['17', 'calculator', 'Very High', 'Interactive form with number inputs, real-time client-side JS calculation, formatted result display. 7 different formulas.', 'Calculator pages'],
    ['18', 'glossary-index', 'Medium', 'Alphabetical A-Z navigation with grouped term links. Supports letter filtering and anchor navigation.', 'Glossary listing page'],
    ['19', 'breadcrumb', 'Low', 'Navigation breadcrumb trail showing page hierarchy path with linked segments.', 'Product pages, Guides, Business types'],
    ['20', 'form', 'High', 'Multi-step form wizard with validation, reCAPTCHA Enterprise, and API submission to /api/signup/v2 and /api/crm/v2.', 'Signup pages'],
    ['21', 'advisor-search', 'Very High', 'Coveo-powered enterprise search with autocomplete, faceted filters, paginated results (617+ advisors), analytics.', 'Advisor directory page'],
    ['22', 'template-preview', 'Medium', 'Template preview card with preview image, download CTA, description, and related templates grid.', 'Template detail pages'],
    ['23', 'video-embed', 'Low', 'Responsive video player embed with play button overlay, thumbnail, and caption.', 'Product feature pages'],
    ['24', 'icon-list', 'Low', 'List items with icon prefix (checkmarks, feature icons) and descriptive text.', 'Product pages, Pricing pages'],
]

create_table(doc, ['#', 'Block', 'Complexity', 'Description & Behaviour', 'Reference URL(s)'], custom_blocks)

doc.add_paragraph()
doc.add_heading('2.3 Block Screenshots', level=2)

add_screenshot(doc, '/tmp/playwright/01-homepage-hero.png', 'Figure 4: Hero Block + Promo Banner + Header', 5.0)
doc.add_paragraph()
add_screenshot(doc, '/tmp/playwright/09-calculator-page.png', 'Figure 5: Calculator Block - Interactive Margin Calculator', 3.5)
doc.add_paragraph()
add_screenshot(doc, '/tmp/playwright/10-compare-page.png', 'Figure 6: Comparison Table Block - Xero vs QuickBooks', 3.5)
doc.add_paragraph()
add_screenshot(doc, '/tmp/playwright/11-advisor-search-page.png', 'Figure 7: Advisor Search Block - Coveo-Powered Search', 3.5)
doc.add_paragraph()
add_screenshot(doc, '/tmp/playwright/13-signup-page.png', 'Figure 8: Form Block - Signup Page with reCAPTCHA', 3.5)

doc.add_page_break()

# SECTION 3: PAGE COUNTS BY TEMPLATE
doc.add_heading('3. Page Counts by Template', level=1)
doc.add_paragraph('The following table summarizes the total page count per template and indicates whether pages can be automatically migrated or require manual intervention.')

page_counts = [
    ['Homepage', '1', 'Manual', 'Most complex page; unique layout with many custom sections'],
    ['Product Feature Page', '24', 'Automated (with QA)', 'Consistent layout; parameterized by feature content'],
    ['Business Type Page', '46', 'Automated', 'Highly templated; identical layout, only content differs per industry'],
    ['Business Type Listing', '2', 'Automated', 'Simple grid layout with linked cards'],
    ['Accountant/Bookkeeper Page', '6', 'Semi-Automated', 'Similar structure but content varies; some manual review needed'],
    ['Practice Product Page', '5', 'Semi-Automated', 'Product-specific layouts; moderate variation'],
    ['Pricing Page', '4', 'Manual', 'Dynamic pricing toggle, plan-specific feature lists, A/B variants'],
    ['Guide Listing Page', '1', 'Automated', 'Simple hub page with card grid'],
    ['Guide Article Page', '37', 'Automated (with QA)', 'Long-form content; consistent article template'],
    ['Template Listing Page', '1', 'Automated', 'Simple hub page with card grid'],
    ['Template Detail Page', '22', 'Automated (with QA)', 'Consistent template; preview + download + description'],
    ['Glossary Listing Page', '1', 'Automated', 'A-Z index; programmatic generation'],
    ['Glossary Term Page', '115', 'Automated', 'Highly templated; identical structure, only definition content varies'],
    ['Calculator Listing Page', '1', 'Automated', 'Simple hub page'],
    ['Calculator Tool Page', '7', 'Manual', 'Interactive JS calculators; each requires custom calculation logic'],
    ['Compare Page', '3', 'Semi-Automated', 'Structured comparison tables; competitor-specific content'],
    ['Company Info Page', '11', 'Semi-Automated', 'Varied corporate content; moderate manual review'],
    ['Legal Page', '3', 'Automated', 'Text-heavy; straightforward content migration'],
    ['Campaign Page', '7', 'Manual', 'Unique promotional layouts; forms and dynamic content'],
    ['Resource Page', '5', 'Semi-Automated', 'Hub pages with curated card grids'],
    ['AI Feature Page', '1', 'Manual', 'Unique AI showcase with demos and interactive elements'],
    ['Utility Page', '6', 'Mixed', 'Advisor Finder (Manual); others Automated'],
    ['Signup Page', '6', 'Manual', 'Multi-step forms with API integration and reCAPTCHA'],
    ['TOTAL', '~313', '', ''],
]

create_table(doc, ['Template', 'Page Count', 'Migration Type', 'Notes'], page_counts)

doc.add_paragraph()
doc.add_heading('3.1 Migration Type Summary', level=2)

migration_summary = [
    ['Automated', '~233', '74%', 'Glossary (115), Business Types (46), Product Features (24), Guides (37), Templates (22), Legal (3), Listings (6)'],
    ['Semi-Automated', '~27', '9%', 'Accountant pages (6), Practice products (5), Compare (3), Company info (11), Resources (5)'],
    ['Manual', '~53', '17%', 'Homepage (1), Pricing (4), Calculators (7), Campaigns (7), AI page (1), Signup (6), Advisors (1)'],
    ['Total', '~313', '100%', ''],
]

create_table(doc, ['Type', 'Pages', '% of Total', 'Includes'], migration_summary)

doc.add_page_break()

# SECTION 4: INTEGRATIONS ANALYSIS
doc.add_heading('4. Integrations Analysis', level=1)
doc.add_paragraph('The Xero.com site relies on 27 third-party services categorized by function and criticality.')

doc.add_heading('4.1 Critical Integrations (Functional)', level=2)

critical_integrations = [
    ['1', 'Xero Signup API', 'API', 'High', '/api/signup/v2 - account registration with plan selection, multi-step form', 'Signup pages (/us/signup/*)'],
    ['2', 'Xero CRM API', 'API', 'Medium', '/api/crm/v2 - Salesforce-style CRM integration for lead capture', 'Signup pages, lead forms'],
    ['3', 'Google reCAPTCHA Enterprise', 'API/Embed', 'Medium', 'Invisible reCAPTCHA protecting signup forms', 'Signup pages'],
    ['4', 'Coveo Search', 'API/Embed', 'Very High', 'xeroprod.org.coveo.com - enterprise search with faceted filters, autocomplete, paginated results', 'Advisor directory (/us/advisors/)'],
    ['5', 'Stripe.js', 'API/Embed', 'Medium', 'js.stripe.com/v3 - payment infrastructure for checkout/payment flows', 'All pages (global)'],
    ['6', 'Intercom', 'Embed', 'Low', 'widget.intercom.io - live chat/support widget', 'All pages (global)'],
    ['7', 'OneTrust', 'Embed', 'Medium', 'cdn-au.onetrust.com - cookie consent, CCPA compliance, preference center', 'All pages (global)'],
]

create_table(doc, ['#', 'Integration', 'Type', 'Complexity', 'Description', 'Reference URL(s)'], critical_integrations)

doc.add_heading('4.2 Analytics & Tracking (Delayed Loading)', level=2)

analytics = [
    ['8', 'Google Tag Manager', 'Embed', 'Low', '2 containers - orchestrates all tracking tags', 'All pages'],
    ['9', 'Google Analytics', 'Embed', 'Low', '3 properties', 'All pages'],
    ['10', 'Google Ads + DoubleClick', 'Embed', 'Low', 'Ad conversion tracking (4 IDs)', 'All pages'],
    ['11', 'Facebook/Meta Pixel', 'Embed', 'Low', 'Social ad tracking', 'All pages'],
    ['12', 'TikTok Pixel', 'Embed', 'Low', '2 pixel IDs', 'All pages'],
    ['13', 'LinkedIn Insight Tag', 'Embed', 'Low', 'B2B ad tracking', 'All pages'],
    ['14', 'Reddit Pixel', 'Embed', 'Low', 'Ad tracking', 'All pages'],
    ['15', 'Bing UET Tag', 'Embed', 'Low', 'Microsoft ad conversion', 'All pages'],
    ['16', 'Amazon Ads', 'Embed', 'Low', 'Ad tracking', 'All pages'],
    ['17', 'Contentsquare', 'Embed', 'Low', 'UX/heatmap analytics', 'All pages'],
    ['18', 'Demandbase', 'Embed', 'Low', 'ABM/company identification', 'All pages'],
    ['19', 'G2 Crowd', 'Embed', 'Low', 'Review site attribution', 'All pages'],
    ['20', 'Adalyser', 'Embed', 'Low', 'Ad measurement', 'All pages'],
    ['21', 'New Relic', 'API', 'Low', 'Application performance monitoring', 'All pages'],
]

create_table(doc, ['#', 'Integration', 'Type', 'Complexity', 'Description', 'Pages'], analytics)

doc.add_heading('4.3 CRO, Content & Infrastructure', level=2)

other_integrations = [
    ['22', 'ConvertFlow', 'Embed', 'Low', 'CRO pop-ups and experiments', 'All pages'],
    ['23', 'SchemaApp', 'Embed', 'Low', 'Structured data / Schema.org markup', 'All pages'],
    ['24', 'OpenAI SDK', 'Embed', 'Low', 'AI chat widget', 'All pages'],
    ['25', 'AWS S3', 'Infrastructure', 'Low', 'Advisor profile images', 'Advisor directory'],
    ['26', 'Xero Events API', 'API', 'Medium', '/api/events/v1/ - custom event tracking', 'All pages'],
    ['27', 'Xero User Tracking', 'API', 'Medium', 'dgt-digital-marketing.xero.com - marketing identification', 'All pages'],
]

create_table(doc, ['#', 'Integration', 'Type', 'Complexity', 'Description', 'Pages'], other_integrations)

doc.add_heading('4.4 Integration Summary', level=2)

int_summary = [
    ['Critical (Functional)', '7', '~15 days'],
    ['Analytics & Tracking', '14', '~5 days (mostly GTM-managed)'],
    ['CRO, Content & Infrastructure', '6', '~5 days'],
    ['Total', '27', '~25 days'],
]

create_table(doc, ['Category', 'Count', 'Estimated Effort'], int_summary)

doc.add_page_break()

# SECTION 5: COMPLEX USE CASES
doc.add_heading('5. Complex Use Cases & Observations', level=1)
doc.add_paragraph('The following complex behaviours and functionality require special attention during migration.')

complex_cases = [
    ['1', 'Coveo-Powered Advisor Search', '1 page', '/us/advisors/', 'Very High', 'Enterprise search with autocomplete, faceted filters (Location, Industry), paginated results (617+ advisors). Requires Coveo SDK and API key management.'],
    ['2', 'Interactive Calculators', '7 pages', '/us/calculators/*', 'High', 'Each calculator has unique formulas (margin, markup, cash flow, income tax, net profit, sales tax, timesheet). All client-side JS, needs vanilla JS rebuild.'],
    ['3', 'Multi-Step Signup Forms', '6 pages', '/us/signup/*', 'High', 'React wizard with conditional logic, reCAPTCHA Enterprise, Salesforce CRM field mapping, dual API submission (/api/signup/v2 + /api/crm/v2).'],
    ['4', 'Dynamic Pricing Toggle', '~8 pages', 'Homepage, Pricing, Products', 'Medium-High', 'Radio toggle switches between Buy Now/Free Trial pricing. Different prices, savings, and signup URLs per variant.'],
    ['5', 'Countdown Timer Banner', 'All pages', 'Global promo banner', 'Medium', 'Live countdown timer to promotion deadline. JS timer logic, dismiss cookie persistence, campaign-configurable offer text.'],
    ['6', 'Mega Navigation Menu', 'All pages', 'Global header', 'Medium', 'Multi-level navigation with product categories, features, resources. Desktop mega-menu and mobile hamburger drawer.'],
    ['7', 'Template Generators', '3 pages', '/us/templates/*-generator/', 'High', 'Interactive tools generating PDFs from user input. May require iframe embed or custom block.'],
    ['8', 'Video Embeds', '~5 pages', 'Product feature pages', 'Low-Medium', 'Custom video player with play button overlay, thumbnail, caption. Lazy loading required.'],
    ['9', 'Region/Locale Selector', 'All pages', 'Footer', 'Medium', 'Switches between US, UK, AU, NZ and 10+ country sites with different pricing, currency, and legal content.'],
    ['10', 'A/B Test Variants', 'Multiple', 'Homepage, Pricing, Signup', 'Medium', 'ConvertFlow experiments. Different signup URLs (early vs early11) suggest active A/B tests on pricing/offers.'],
]

create_table(doc, ['#', 'Use Case', 'Instances', 'Location', 'Complexity', 'Description'], complex_cases)

doc.add_page_break()

# SECTION 6: MIGRATION ESTIMATES
doc.add_heading('6. Migration Estimates', level=1)

doc.add_heading('6.1 Effort Breakdown by Phase', level=2)

effort_phases = [
    ['Phase 1: Foundation', '2 weeks', 'Design system extraction, global header with mega-menu, global footer with region selector, promo banner, cookie consent'],
    ['Phase 2: Core Templates', '4 weeks', 'Homepage, Product Features (24), Pricing (4), Business Types (46), Compare (3). Build hero, cards, columns, pricing-table, comparison-table, tabs, accordion blocks'],
    ['Phase 3: Content Templates', '3 weeks', 'Guides (37), Templates (22), Glossary (115), Calculators (7), Resources (5). Build calculator, glossary-index, template-preview blocks'],
    ['Phase 4: Specialized Pages', '2 weeks', 'Signup (6) with form + API, Campaign (7), AI Feature page, Company/Legal, Utility, Advisor Search (Coveo)'],
    ['Phase 5: Bulk Import & QA', '2 weeks', 'Import all 313 pages, content QA, responsive testing, cross-browser validation'],
    ['Phase 6: Polish & Launch', '1 week', 'Accessibility audit (WCAG 2.1 AA), SEO validation, performance optimization (Lighthouse 100), URL redirect mapping'],
]

create_table(doc, ['Phase', 'Duration', 'Details'], effort_phases)

doc.add_heading('6.2 Effort by Work Type', level=2)

effort_type = [
    ['Block Development (6 core + 18 custom)', '53 days', '~10.5 weeks'],
    ['Template & Import Infrastructure', '30 days', '~6 weeks'],
    ['Integrations (27 total)', '25 days', '~5 weeks'],
    ['Design System & Global Styles', '10 days', '~2 weeks'],
    ['Content Migration & QA', '18 days', '~3.5 weeks'],
    ['Performance & Accessibility', '8 days', '~1.5 weeks'],
    ['Total', '~144 days', '~29 person-weeks'],
]

create_table(doc, ['Work Type', 'Effort (Days)', 'Duration'], effort_type)

doc.add_heading('6.3 Resource Recommendation', level=2)

resources = [
    ['EDS Developer', '2', 'Full (14 weeks)', 'Block development, import infrastructure, integrations'],
    ['Content/Import Engineer', '1', 'Weeks 4-12', 'Bulk import, content QA, validation'],
    ['Design/CSS Specialist', '1', 'Weeks 1-6', 'Design system extraction, visual fidelity, responsive polish'],
    ['QA/Accessibility Tester', '1', 'Weeks 10-14', 'Cross-browser testing, accessibility audit, performance'],
    ['Project Manager', '1', 'Full', 'Coordination, stakeholder communication'],
]

create_table(doc, ['Role', 'Count', 'Duration', 'Responsibilities'], resources)

doc.add_heading('6.4 Timeline Summary', level=2)

timeline = [
    ['Total Pages', '~313'],
    ['Unique Templates', '23'],
    ['Unique Blocks', '24 (6 core + 18 custom)'],
    ['Integrations', '27'],
    ['Total Effort', '~144 person-days (~29 person-weeks)'],
    ['Calendar Duration', '~14 weeks (with team of 4-5)'],
    ['Automated Migration', '233 pages (74%)'],
    ['Manual Migration', '53 pages (17%)'],
    ['Semi-Automated', '27 pages (9%)'],
]

create_table(doc, ['Metric', 'Value'], timeline)

doc.add_heading('6.5 Risk Factors', level=2)

risks = [
    ['Coveo Advisor Search', 'Very High', 'Enterprise search SDK; requires API keys and faceted config', 'Prototype early; consider static listing as fallback'],
    ['Interactive Calculators (7)', 'High', 'Custom vanilla JS logic; currently React-based', 'Build generic calculator block framework'],
    ['Signup Forms + API', 'High', 'Backend API integration with reCAPTCHA', 'Coordinate with Xero backend team'],
    ['Template Generators (3)', 'High', 'PDF generation from user input', 'May require iframe embed approach'],
    ['A/B Test Variants', 'Medium', 'Multiple pricing/offer variants active', 'Map all variant URLs; coordinate with marketing'],
    ['115 Glossary Pages', 'Low risk, High volume', 'Any template bug multiplies 115x', 'Thorough QA before bulk import'],
]

create_table(doc, ['Risk Area', 'Impact', 'Description', 'Mitigation'], risks)

doc.add_page_break()

# APPENDIX
doc.add_heading('Appendix: Additional Template Screenshots', level=1)

screenshots = [
    ('/tmp/playwright/05-business-type-page.png', 'Business Type Page - Construction'),
    ('/tmp/playwright/06-guide-article-page.png', 'Guide Article Page - Invoicing Guide'),
    ('/tmp/playwright/07-glossary-term-page.png', 'Glossary Term Page - Cash Flow'),
    ('/tmp/playwright/08-template-detail-page.png', 'Template Detail Page - Invoice Template'),
    ('/tmp/playwright/12-company-page.png', 'Company Info Page - About Xero'),
]

for path, caption in screenshots:
    add_screenshot(doc, path, caption, 3.5)
    doc.add_paragraph()

output_path = '/workspace/migration-work/Xero-EDS-Migration-Analysis.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')
print(f'Size: {os.path.getsize(output_path) / 1024 / 1024:.1f} MB')
